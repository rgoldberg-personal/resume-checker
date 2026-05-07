"""DOCX text extraction using python-docx."""
from __future__ import annotations

import docx


def extract_docx_text(file_path: str) -> dict:
    """
    Extract text from a DOCX file using python-docx.

    Extracts from both paragraphs and tables to handle table-based CV layouts.

    Returns:
        {
            "text": str,
            "paragraph_count": int,
        }
    """
    document = docx.Document(file_path)

    parts: list[str] = []

    # Extract from paragraphs
    for p in document.paragraphs:
        if p.text.strip():
            parts.append(p.text)

    # Extract from tables (many CVs use table-based layouts)
    for table in document.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                parts.append(" | ".join(row_texts))

    text = "\n".join(parts).strip()

    return {
        "text": text,
        "paragraph_count": len(parts),
    }
